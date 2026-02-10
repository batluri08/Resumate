"""
Document Parser Service - Extracts content from PDF and DOCX files
while preserving structure information for later reconstruction
"""

from docx import Document
from PyPDF2 import PdfReader
import os
from typing import Tuple, Dict, List, Any


class DocumentParser:
    """Parse PDF and DOCX documents to extract content and structure"""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a document and return its content and structure
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (text_content, structure_info)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a DOCX file preserving structure information
        """
        doc = Document(file_path)
        
        content_parts = []
        structure = {
            "type": "docx",
            "paragraphs": [],
            "tables": [],
            "sections": []
        }
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                content_parts.append(text)
                
                # Store paragraph structure info
                para_info = {
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else None,
                    "runs": []
                }
                
                # Store run-level formatting
                for run in para.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name if run.font else None,
                        "font_size": str(run.font.size) if run.font and run.font.size else None
                    }
                    para_info["runs"].append(run_info)
                
                structure["paragraphs"].append(para_info)
        
        # Parse tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                content_parts.extend([cell for cell in row_data if cell])
            
            structure["tables"].append({
                "index": i,
                "data": table_data
            })
        
        full_content = "\n".join(content_parts)
        return full_content, structure
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a PDF file - note that PDF formatting is harder to preserve
        """
        reader = PdfReader(file_path)
        
        content_parts = []
        structure = {
            "type": "pdf",
            "pages": [],
            "total_pages": len(reader.pages)
        }
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                content_parts.append(text)
                structure["pages"].append({
                    "index": i,
                    "text": text
                })
        
        full_content = "\n".join(content_parts)
        return full_content, structure
    
    def extract_sections(self, content: str) -> Dict[str, str]:
        """
        Extract common resume sections from content
        """
        sections = {}
        
        # Common section headers in resumes
        section_keywords = [
            "summary", "objective", "profile",
            "experience", "work experience", "employment",
            "education", "academic",
            "skills", "technical skills", "competencies",
            "projects", "achievements", "accomplishments",
            "certifications", "certificates",
            "languages", "interests", "hobbies",
            "references", "publications"
        ]
        
        lines = content.split("\n")
        current_section = "header"
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            is_header = False
            for keyword in section_keywords:
                if keyword in line_lower and len(line_lower) < 50:
                    # Save previous section
                    if current_content:
                        sections[current_section] = "\n".join(current_content)
                    
                    current_section = keyword
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections
