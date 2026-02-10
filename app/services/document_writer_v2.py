"""
Document Writer Service - Applies find/replace changes while preserving formatting
"""

from docx import Document
from docx.shared import Pt
import os
from typing import List, Dict
from pdf2docx import Converter
import tempfile
import shutil
import re


class DocumentWriter:
    """Apply find/replace changes to documents while preserving formatting"""
    
    def write(
        self,
        original_path: str,
        output_path: str,
        changes: List[Dict],
        file_type: str
    ) -> str:
        """Apply changes to document."""
        if file_type == ".pdf":
            docx_output = output_path.replace('.pdf', '.docx')
            return self._write_pdf_to_docx(original_path, docx_output, changes)
        else:
            return self._write_docx(original_path, output_path, changes)
    
    def _write_pdf_to_docx(
        self,
        original_path: str,
        output_path: str,
        changes: List[Dict]
    ) -> str:
        """Convert PDF to DOCX and apply changes."""
        temp_dir = tempfile.mkdtemp()
        
        try:
            temp_docx = os.path.join(temp_dir, "converted.docx")
            cv = Converter(original_path)
            cv.convert(temp_docx)
            cv.close()
            
            self._write_docx(temp_docx, output_path, changes)
            return output_path
            
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    def _write_docx(
        self,
        original_path: str,
        output_path: str,
        changes: List[Dict]
    ) -> str:
        """Apply find/replace changes to DOCX preserving formatting."""
        shutil.copy2(original_path, output_path)
        doc = Document(output_path)
        
        print(f"[DEBUG] Applying {len(changes)} changes to document")
        
        total_replacements = 0
        
        for i, change in enumerate(changes):
            find_text = change.get("find", "")
            replace_text = change.get("replace", "")
            
            if not find_text or not replace_text:
                continue
            
            # Allow replacement to be up to 50 characters longer for better optimization
            # Resume improvements often need more words for keywords and details
            diff = len(replace_text) - len(find_text)
            if diff > 50:
                print(f"[DEBUG] Change {i+1}: SKIPPED (too long +{diff} chars) - '{find_text[:40]}...'")
                continue
            
            # Skip if find text is too short (likely a fragment) - reduced to 15 chars
            if len(find_text) < 15:
                print(f"[DEBUG] Change {i+1}: SKIPPED (too short/fragment) - '{find_text}'")
                continue
            
            # Ensure bullet point preservation
            replace_text = self._preserve_bullet_format(find_text, replace_text)
            
            # Apply to all paragraphs (document body and tables)
            all_paras = self._get_all_paragraphs(doc)
            count = 0
            
            for para in all_paras:
                if find_text in para.text:
                    print(f"[DEBUG] FOUND in para: '{para.text[:60]}...'")
                    self._smart_replace(para, find_text, replace_text)
                    count += 1
            
            if count > 0:
                total_replacements += count
                print(f"[DEBUG] Change {i+1}: Applied {count}x - '{find_text[:50]}...'")
            else:
                # Try normalized match (handles whitespace differences)
                norm_count = self._apply_normalized_change(all_paras, find_text, replace_text)
                if norm_count > 0:
                    total_replacements += norm_count
                    print(f"[DEBUG] Change {i+1}: Normalized match {norm_count}x - '{find_text[:50]}...'")
                else:
                    # Log what we're looking for vs what's in doc
                    print(f"[DEBUG] Change {i+1}: NOT FOUND - '{find_text}'")
                    # Try to find similar text
                    for para in all_paras[:10]:
                        if len(para.text) > 20 and find_text[:10].lower() in para.text.lower():
                            print(f"[DEBUG]   Similar para found: '{para.text[:80]}...'")
        
        print(f"[DEBUG] Total changes applied: {total_replacements}")
        
        doc.save(output_path)
        return output_path
    
    def _preserve_bullet_format(self, original: str, replacement: str) -> str:
        """Ensure replacement preserves bullet point formatting from original."""
        bullet_chars = ['•', '‣', '⁃', '●', '○', '■', '□', '-', '–', '—', '*']
        
        # Check if original starts with a bullet
        original_stripped = original.lstrip()
        original_indent = original[:len(original) - len(original_stripped)]
        
        for bullet in bullet_chars:
            if original_stripped.startswith(bullet):
                # Original has a bullet - ensure replacement does too
                replacement_stripped = replacement.lstrip()
                replacement_indent = replacement[:len(replacement) - len(replacement_stripped)]
                
                # Check if replacement already has a bullet
                has_bullet = any(replacement_stripped.startswith(b) for b in bullet_chars)
                
                if not has_bullet:
                    # Add the original's bullet and spacing
                    bullet_and_space = original_stripped[:2] if len(original_stripped) > 1 and original_stripped[1] == ' ' else bullet
                    return original_indent + bullet_and_space + replacement_stripped
                else:
                    # Preserve original indentation
                    return original_indent + replacement_stripped
        
        return replacement
    
    def _get_all_paragraphs(self, doc):
        """Get all paragraphs from document including tables."""
        paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras.extend(cell.paragraphs)
        return paras
    
    def _smart_replace(self, para, find_text: str, replace_text: str):
        """
        Replace text in paragraph while preserving formatting structure.
        Key: Keep run boundaries intact, only modify text content.
        """
        full_text = para.text
        
        if find_text not in full_text:
            return
        
        # Simple case: paragraph has no runs or just one run
        if not para.runs or len(para.runs) == 1:
            new_text = full_text.replace(find_text, replace_text, 1)
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.text = new_text
            return
        
        # Complex case: multiple runs
        # Strategy: Find which runs contain the text and do minimal changes
        
        # Build a map of character positions to runs
        run_boundaries = []
        pos = 0
        for run in para.runs:
            run_boundaries.append((pos, pos + len(run.text), run))
            pos += len(run.text)
        
        # Find where the text to replace is
        start_pos = full_text.find(find_text)
        end_pos = start_pos + len(find_text)
        
        # Find which runs are affected
        affected_runs = []
        for run_start, run_end, run in run_boundaries:
            if run_start < end_pos and run_end > start_pos:
                affected_runs.append((run_start, run_end, run))
        
        if not affected_runs:
            return
        
        # If only one run is affected, simple replace within that run
        if len(affected_runs) == 1:
            run_start, run_end, run = affected_runs[0]
            local_start = start_pos - run_start
            local_end = end_pos - run_start
            run.text = run.text[:local_start] + replace_text + run.text[local_end:]
            return
        
        # Multiple runs affected - put replacement in first affected run, clear others
        first_run_start, first_run_end, first_run = affected_runs[0]
        last_run_start, last_run_end, last_run = affected_runs[-1]
        
        # Text before the find in the first run
        prefix = first_run.text[:start_pos - first_run_start]
        # Text after the find in the last run
        suffix = last_run.text[end_pos - last_run_start:]
        
        # Set first run to prefix + replacement + suffix
        first_run.text = prefix + replace_text + suffix
        
        # Clear runs in between (not first, not last if different)
        for run_start, run_end, run in affected_runs[1:]:
            run.text = ""
    
    def _apply_normalized_change(self, paras, find_text: str, replace_text: str) -> int:
        """Try to match with normalized whitespace."""
        count = 0
        find_norm = self._normalize(find_text)
        
        for para in paras:
            para_norm = self._normalize(para.text)
            
            if find_norm in para_norm:
                # Found normalized match - try to find actual substring
                # Use a more flexible match
                pattern = self._create_flexible_pattern(find_text)
                match = re.search(pattern, para.text, re.IGNORECASE)
                
                if match:
                    actual_find = match.group(0)
                    self._smart_replace(para, actual_find, replace_text)
                    count += 1
        
        return count
    
    def _normalize(self, text: str) -> str:
        """Normalize text for comparison."""
        text = re.sub(r'\s+', ' ', text)
        return text.strip().lower()
    
    def _create_flexible_pattern(self, text: str) -> str:
        """Create a regex pattern that allows flexible whitespace."""
        # Escape special regex chars but allow flexible whitespace
        words = text.split()
        escaped = [re.escape(w) for w in words]
        return r'\s*'.join(escaped)
